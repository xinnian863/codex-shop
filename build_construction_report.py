from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("work_docs/郑州市惠济区实验小学教学楼施工组织设计报告-附件2优化版.docx")
SITE_PLAN = Path("work_docs/site_plan_attachment2.png")


def set_run_font(run, east="宋体", west="Times New Roman", size=None, bold=None):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_text(style, east="宋体", west="Times New Roman", size=12, bold=False, color=None):
    style.font.name = west
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    style._element.rPr.rFonts.set(qn("w:ascii"), west)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C1CC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_tbl_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(indent_dxa))
    ind.set(qn("w:type"), "dxa")


def add_para(doc, text="", style=None, align=None, first_line=True, after=6):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=12)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def fill_cell(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_tbl_width(table)
    set_cell_margins(table)
    set_table_borders(table)
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        fill_cell(cell, h, bold=True, size=10)
        set_cell_shading(cell, "F2F4F7")
        if widths:
            set_cell_width(cell, widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(val)) > 10 else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cells[idx], val, size=10, align=align)
            if widths:
                set_cell_width(cells[idx], widths[idx])
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=12)


def draw_site_plan(path: Path):
    from PIL import Image, ImageDraw, ImageFont

    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 1050
    img = Image.new("RGB", (w, h), "#F7F4EC")
    d = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/arial.ttf",
        ]
        for p in candidates:
            try:
                return ImageFont.truetype(p, size)
            except OSError:
                continue
        return ImageFont.load_default()

    f_title = font(38, True)
    f_h = font(25, True)
    f = font(22)
    f_s = font(18)
    f_xs = font(15)

    def center_text(box, text, fill="#1F2933", ft=None):
        ft = ft or f
        lines = text.split("\n")
        heights = []
        widths = []
        for line in lines:
            bb = d.textbbox((0, 0), line, font=ft)
            widths.append(bb[2] - bb[0])
            heights.append(bb[3] - bb[1])
        total_h = sum(heights) + (len(lines) - 1) * 6
        y = box[1] + (box[3] - box[1] - total_h) / 2
        for line, tw, th in zip(lines, widths, heights):
            x = box[0] + (box[2] - box[0] - tw) / 2
            d.text((x, y), line, font=ft, fill=fill)
            y += th + 6

    def rect(box, fill, outline="#38434F", width_line=3, label=None, ft=None):
        d.rounded_rectangle(box, radius=10, fill=fill, outline=outline, width=width_line)
        if label:
            center_text(box, label, ft=ft or f)

    # Sheet frame and title block.
    d.rectangle((34, 34, w - 34, h - 34), outline="#111827", width=4)
    d.line((34, 100, w - 34, 100), fill="#111827", width=3)
    d.text((62, 48), "附件2 施工现场总平面布置示意图", font=f_title, fill="#111827")
    d.text((1260, 58), "工程名称：郑州市惠济区实验小学教学楼", font=f_s, fill="#374151")

    # North arrow.
    d.polygon([(1620, 145), (1588, 225), (1620, 205), (1652, 225)], fill="#111827")
    d.line((1620, 205, 1620, 300), fill="#111827", width=5)
    d.text((1606, 112), "北", font=f_h, fill="#111827")

    # Roads and construction site.
    d.rectangle((95, 132, 1510, 935), fill="#DCE7D2", outline="#2F5138", width=4)
    d.rectangle((95, 132, 1510, 205), fill="#C7D8EA", outline="#2E4E6B", width=2)
    center_text((95, 132, 1510, 205), "城市支路  /  主出入口  /  车辆冲洗平台", fill="#1E3A5F", ft=f_h)
    d.rectangle((135, 225, 1470, 890), outline="#828B57", width=16)
    d.text((1180, 846), "场内临时道路宽4m", font=f_s, fill="#4B5563")

    # Main gate and wash bay.
    rect((650, 205, 790, 255), "#F9FAFB", "#1F2937", 3, "主出入口", f_s)
    rect((805, 213, 960, 252), "#E0F2FE", "#2563EB", 2, "洗车槽", f_s)
    rect((980, 213, 1100, 252), "#F3F4F6", "#4B5563", 2, "门卫", f_s)

    # Building and tower crane.
    rect((520, 410, 1040, 675), "#FDEBD3", "#B45309", 4, "拟建教学楼主体\n5层框架结构", f_h)
    d.rectangle((550, 440, 1010, 645), outline="#EAB308", width=2)
    d.ellipse((420, 320, 1140, 780), outline="#DC2626", width=4)
    d.line((780, 545, 1160, 360), fill="#DC2626", width=5)
    d.ellipse((742, 507, 818, 583), fill="#FEE2E2", outline="#DC2626", width=4)
    center_text((742, 507, 818, 583), "塔吊", fill="#991B1B", ft=f_s)
    d.rectangle((920, 330, 1152, 365), fill="#F7F4EC")
    d.text((932, 336), "QTZ63塔吊覆盖范围", font=f_s, fill="#991B1B")

    # Functional areas.
    rect((175, 250, 405, 345), "#EDE9FE", "#6D28D9", 3, "办公生活区", f)
    rect((175, 375, 405, 500), "#FEF3C7", "#B45309", 3, "钢筋加工棚\n原材/半成品", f_s)
    rect((175, 530, 405, 650), "#FAE8FF", "#A21CAF", 3, "模板加工棚\n木方模板", f_s)
    rect((175, 705, 405, 830), "#E5E7EB", "#4B5563", 3, "周转材料堆场\n钢管扣件", f_s)

    rect((1135, 270, 1420, 365), "#DCFCE7", "#15803D", 3, "砌块堆场", f)
    rect((1135, 400, 1420, 505), "#F3E8FF", "#7E22CE", 3, "砂浆搅拌区", f)
    rect((1135, 540, 1420, 650), "#DBEAFE", "#1D4ED8", 3, "水泥库 / 材料库", f_s)
    rect((1135, 700, 1420, 815), "#E0F2FE", "#0369A1", 3, "沉淀池\n消防水源", f_s)

    # Utilities.
    d.line((150, 865, 1450, 865), fill="#2563EB", width=4)
    d.text((152, 875), "临水及消防管线", font=f_xs, fill="#1D4ED8")
    for x in range(180, 1450, 90):
        d.line((x, 865, x + 28, 850), fill="#2563EB", width=2)
    d.line((135, 235, 135, 845, 1450, 845), fill="#F59E0B", width=4)
    d.text((142, 810), "临电线路沿围挡敷设，三级配电二级保护", font=f_xs, fill="#92400E")
    for x, y in [(420, 845), (760, 845), (1110, 845), (135, 560), (135, 330)]:
        d.rectangle((x - 12, y - 12, x + 12, y + 12), fill="#FEF3C7", outline="#92400E", width=2)

    # Safety notes.
    d.rectangle((470, 725, 1085, 830), fill="#FFFFFF", outline="#6B7280", width=2)
    d.text((492, 740), "安全文明控制：封闭围挡、喷淋降尘、车辆冲洗、材料分区堆放。", font=f_xs, fill="#374151")
    d.text((492, 772), "消防器材定点配置；材料车辆避开上下学时段，主通道保持畅通。", font=f_xs, fill="#374151")
    d.text((492, 804), "塔吊回转半径内设置警戒标识。", font=f_xs, fill="#374151")

    # Legend and title block.
    legend = (1535, 330, 1748, 735)
    d.rectangle(legend, fill="#FFFFFF", outline="#111827", width=3)
    d.text((1560, 350), "图例", font=f_h, fill="#111827")
    items = [
        ("#FDEBD3", "拟建建筑物"),
        ("#FEF3C7", "加工及材料区"),
        ("#EDE9FE", "办公生活区"),
        ("#C7D8EA", "场外道路"),
        ("#DC2626", "塔吊覆盖"),
        ("#2563EB", "临水管线"),
        ("#F59E0B", "临电线路"),
    ]
    yy = 405
    for color, label in items:
        if color in ("#DC2626", "#2563EB", "#F59E0B"):
            d.line((1562, yy + 12, 1605, yy + 12), fill=color, width=5)
        else:
            d.rectangle((1562, yy, 1605, yy + 24), fill=color, outline="#374151", width=1)
        d.text((1620, yy - 2), label, font=f_xs, fill="#374151")
        yy += 42

    title_box = (1230, 895, 1748, 1008)
    d.rectangle(title_box, fill="#FFFFFF", outline="#111827", width=3)
    rows = [895, 933, 971, 1008]
    for y in rows[1:-1]:
        d.line((1230, y, 1748, y), fill="#111827", width=2)
    d.line((1400, 895, 1400, 1008), fill="#111827", width=2)
    d.text((1244, 905), "绘图单位", font=f_xs, fill="#111827")
    d.text((1415, 905), "管理学院工程管理系", font=f_xs, fill="#111827")
    d.text((1244, 943), "绘图人", font=f_xs, fill="#111827")
    d.text((1415, 943), "闫恒", font=f_xs, fill="#111827")
    d.text((1244, 981), "审核人", font=f_xs, fill="#111827")
    d.text((1415, 981), "指导教师", font=f_xs, fill="#111827")

    img.save(path, quality=95)
    return path


def configure_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    styles = doc.styles
    style_text(styles["Normal"], size=12)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(6)
    style_text(styles["Heading 1"], size=14, bold=True, color="000000")
    style_text(styles["Heading 2"], size=12, bold=True, color="000000")
    style_text(styles["Heading 3"], size=12, bold=True, color="000000")
    style_text(styles["List Bullet"], size=12)
    style_text(styles["List Number"], size=12)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("郑州市惠济区实验小学教学楼施工组织设计")
    set_run_font(r, size=9)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("工程管理专业")
    set_run_font(r, east="黑体", size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《建筑工程施工组织与进度控制》")
    set_run_font(r, east="黑体", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("课程设计报告")
    set_run_font(r, east="黑体", size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（2025-2026学年第二学期）")
    set_run_font(r, size=14)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("郑州市惠济区实验小学教学楼施工组织设计")
    set_run_font(r, east="黑体", size=18, bold=True)

    for _ in range(3):
        doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    set_cell_margins(table, top=100, bottom=100)
    rows = [
        ("班    级：", "2024级工程管理X班"),
        ("学    号：", "542413430627"),
        ("姓    名：", "闫恒"),
        ("指导教师：", "按系统"),
        ("成    绩：", ""),
    ]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            fill_cell(table.rows[i].cells[j], val, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(table.rows[i].cells[j], 1800 if j == 0 else 4200)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("管理学院  工程管理系")
    set_run_font(r, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("二〇二六年六月")
    set_run_font(r, size=12)
    doc.add_page_break()


def toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, east="黑体", size=15, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中自动更新。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    doc.add_page_break()


def body(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("郑州市惠济区实验小学教学楼施工组织设计")
    set_run_font(r, east="黑体", size=16, bold=True)

    add_heading(doc, "1 编制依据", 1)
    add_heading(doc, "1.1 施工合同及设计文件", 2)
    add_para(doc, "本施工组织设计以郑州市惠济区实验小学教学楼工程施工总承包任务为对象，依据招标文件、施工合同、施工图纸、岩土工程勘察报告及现场踏勘资料编制。文件用于指导本单位工程施工准备、施工部署、资源组织、进度控制、质量安全管理和现场平面布置。")
    add_table(
        doc,
        ["序号", "文件名称", "主要作用"],
        [
            ["1", "郑州市惠济区实验小学教学楼工程施工图", "确定建筑、结构、装饰及专业配合施工内容"],
            ["2", "郑州市惠济区实验小学教学楼工程施工合同", "明确合同范围、工期、质量、安全和结算要求"],
            ["3", "岩土工程勘察报告", "确定地基基础施工条件和地下水、土层参数"],
            ["4", "施工现场踏勘记录", "确定场地道路、临水临电、周边环境和材料运输条件"],
            ["5", "工程预算书及主要工程量清单", "确定劳动力、材料、机械和资金安排"],
        ],
        widths=[900, 3600, 4860],
    )
    add_heading(doc, "1.2 法律法规、标准规范及施工范围", 2)
    add_para(doc, "本工程执行国家现行法律法规、工程建设强制性标准及河南省、郑州市有关建筑施工管理规定。施工过程坚持依法施工、按图施工、按规范验收，并将质量、安全、环保和文明施工要求落实到各分部分项工程。")
    add_table(
        doc,
        ["类别", "名称", "编号"],
        [
            ["法律法规", "中华人民共和国建筑法", "主席令第91号"],
            ["法律法规", "建设工程质量管理条例", "国务院令第279号"],
            ["法律法规", "建设工程安全生产管理条例", "国务院令第393号"],
            ["验收规范", "建筑工程施工质量验收统一标准", "GB 50300-2013"],
            ["验收规范", "混凝土结构工程施工质量验收规范", "GB 50204-2015"],
            ["验收规范", "砌体结构工程施工质量验收规范", "GB 50203-2011"],
            ["安全规范", "建筑施工安全检查标准", "JGJ 59-2011"],
            ["技术规范", "建筑施工扣件式钢管脚手架安全技术规范", "JGJ 130-2011"],
            ["技术规范", "建筑施工模板安全技术规范", "JGJ 162-2008"],
            ["环保文明", "建筑施工现场环境与卫生标准", "JGJ 146-2013"],
        ],
        widths=[1600, 5600, 2160],
    )
    add_para(doc, "施工范围包括土方开挖、地基与基础、主体钢筋混凝土结构、砌体围护、屋面防水、楼地面、内外墙装饰、门窗、室外散水台阶及与土建施工相关的预留预埋和专业配合。安装工程按土建施工进度进行穿插配合，本课程设计重点考虑土建部分施工组织。")

    add_heading(doc, "2 工程概况", 1)
    add_heading(doc, "2.1 工程建设概况", 2)
    add_caption(doc, "表2.1 项目基本情况")
    add_table(
        doc,
        ["序号", "项目", "内容"],
        [
            ["1", "工程名称", "郑州市惠济区实验小学教学楼工程"],
            ["2", "建设单位", "郑州市惠济区教育局"],
            ["3", "勘测单位", "河南省建筑设计研究院岩土工程分院"],
            ["4", "设计单位", "郑州市建筑设计院有限公司"],
            ["5", "监理单位", "河南中建工程监理有限公司"],
            ["6", "施工单位", "河南省第一建筑工程集团有限公司"],
            ["7", "建筑面积", "约12680 m2"],
            ["8", "计划工期", "300日历天"],
            ["9", "质量目标", "合格，争创郑州市优质结构工程"],
        ],
        widths=[900, 2200, 6260],
    )
    add_heading(doc, "2.2 建筑设计概况", 2)
    add_para(doc, "本工程为小学教学楼，地上五层，局部一层连廊，建筑高度23.40m。建筑平面呈近似矩形布置，主要功能包括普通教室、教师办公室、多功能教室、卫生间、楼梯间和设备用房。建筑耐火等级为二级，屋面防水等级为一级，设计使用年限50年。")
    add_caption(doc, "表2.2 建筑设计概况")
    add_table(
        doc,
        ["项目", "内容", "项目", "内容"],
        [
            ["总建筑面积", "12680 m2", "占地面积", "约2580 m2"],
            ["建筑功能", "小学教学楼", "建筑高度", "23.40 m"],
            ["层数", "地上5层", "层高", "首层4.20 m，标准层3.90 m"],
            ["外墙", "真石漆外墙，局部干挂装饰板", "楼地面", "水磨石、防滑地砖、局部PVC地面"],
            ["内墙", "乳胶漆墙面，卫生间釉面砖", "顶棚", "乳胶漆顶棚，走廊局部矿棉板吊顶"],
            ["屋面", "保温屋面，SBS改性沥青防水卷材", "门窗", "断桥铝合金窗、防火门、木门"],
        ],
        widths=[1900, 2780, 1900, 2780],
    )
    add_heading(doc, "2.3 结构设计概况", 2)
    add_para(doc, "本工程结构形式为现浇钢筋混凝土框架结构，基础形式采用柱下独立基础局部条形基础。结构安全等级二级，抗震设防烈度7度，框架抗震等级三级。混凝土强度等级基础垫层C15，基础C30，主体梁板柱C30，构造柱、圈梁C25。")
    add_caption(doc, "表2.3 结构设计概况")
    add_table(
        doc,
        ["部位", "设计内容", "主要参数"],
        [
            ["地基基础", "柱下独立基础、局部条形基础", "持力层为粉质黏土，地基承载力特征值约180kPa"],
            ["主体结构", "现浇钢筋混凝土框架结构", "柱截面多为600x600、500x500；梁截面多为250x600、300x650"],
            ["楼屋面板", "现浇钢筋混凝土板", "板厚100-120mm，屋面板局部120mm"],
            ["砌体", "加气混凝土砌块填充墙", "外墙200mm，内隔墙100/200mm"],
            ["抗震", "框架抗震等级三级", "设防烈度7度，设计基本地震加速度0.15g"],
        ],
        widths=[1700, 3300, 4360],
    )
    add_heading(doc, "2.4 周边环境与施工条件", 2)
    add_para(doc, "拟建场地位于郑州市惠济区已建校园内，北侧为城市支路，东侧为现状操场，南侧为已建教学辅助用房，西侧为临时围墙及校区出入口。场地较为平整，施工期间需重点控制噪声、扬尘、材料运输和学生通行安全。现场临水、临电由校区北侧接入，施工道路可沿场地西侧布置环形运输通道。")
    add_table(
        doc,
        ["项目", "条件说明"],
        [
            ["现场地形", "场地基本平整，局部需清表和硬化，满足临建和材料堆场布置要求。"],
            ["周边建筑", "南侧邻近既有建筑，施工时应设置封闭围挡和安全通道。"],
            ["交通运输", "北侧城市支路可供材料车辆进出，运输高峰应避开上下学时段。"],
            ["气象条件", "郑州地区夏季高温多雨，冬季低温干燥，需安排雨季和冬季施工措施。"],
            ["临水临电", "从校区既有管网和配电箱接入，现场设置二级配电和消防水源。"],
            ["临设布置", "办公生活区布置于场地西北角，钢筋加工区、木工加工区靠近塔吊覆盖范围。"],
        ],
        widths=[2200, 7160],
    )
    add_heading(doc, "2.5 主要分项工程量", 2)
    add_para(doc, "根据建筑面积、结构形式及类似工程经验估算主要工程量，作为施工资源配置和进度计划编制依据。实际施工时应以施工图预算和现场签证计量为准。")
    add_table(
        doc,
        ["序号", "分项工程", "单位", "估算工程量", "说明"],
        [
            ["1", "土方开挖", "m3", "8200", "含基础及局部管沟土方"],
            ["2", "基础混凝土", "m3", "1650", "含垫层、独立基础、基础梁"],
            ["3", "主体混凝土", "m3", "4950", "柱、梁、板、楼梯"],
            ["4", "钢筋工程", "t", "780", "基础及主体结构钢筋"],
            ["5", "模板工程", "m2", "38200", "基础、柱、梁、板模板"],
            ["6", "砌体工程", "m3", "2450", "填充墙及局部砌体"],
            ["7", "屋面防水", "m2", "2650", "保温及卷材防水"],
            ["8", "内墙抹灰", "m2", "28600", "教室、走廊及辅助房间"],
            ["9", "楼地面工程", "m2", "11900", "水磨石、地砖及其他面层"],
            ["10", "外墙装饰", "m2", "7200", "保温、抗裂砂浆、真石漆"],
        ],
        widths=[800, 2100, 900, 1500, 4060],
    )

    add_heading(doc, "3 施工部署", 1)
    add_heading(doc, "3.1 施工目标", 2)
    add_para(doc, "本工程施工部署坚持“先地下、后地上；先主体、后围护；先结构、后装饰；安装预留预埋穿插进行”的原则。通过合理划分施工段、组织流水施工和过程控制，保证合同工期、质量、安全、文明施工和成本目标实现。")
    add_bullets(
        doc,
        [
            "工期目标：总工期300日历天，基础工程控制在60日历天内完成，主体结构控制在150日历天内封顶，装饰装修和室外收尾在90日历天内完成。",
            "质量目标：单位工程一次验收合格，主体结构观感质量良好，资料同步、完整、真实。",
            "安全目标：杜绝较大及以上安全事故，轻伤事故频率控制在3‰以内，现场安全检查达到合格以上标准。",
            "环境目标：施工扬尘、噪声、污水和固体废弃物排放满足郑州市文明施工管理要求。",
            "成本目标：通过材料计划、限额领料、机械利用率控制和二次搬运减少，降低非生产性消耗。",
        ],
    )
    add_heading(doc, "3.2 项目管理机构与岗位职责", 2)
    add_para(doc, "项目部采用项目经理负责制，下设施工技术、质量安全、材料设备、预算资料和综合后勤等岗位。各岗位职责清晰、分工协作，形成计划、实施、检查、改进的闭环管理。")
    add_table(
        doc,
        ["岗位", "主要职责"],
        [
            ["项目经理", "全面负责工程进度、质量、安全、成本和外部协调，是项目施工第一责任人。"],
            ["技术负责人", "负责施工组织设计、专项施工方案、技术交底、图纸会审和技术复核。"],
            ["施工员", "负责现场施工安排、工序协调、劳动力调配和施工日志记录。"],
            ["质量员", "负责检验批、分项分部工程质量检查、隐蔽验收和质量整改闭合。"],
            ["安全员", "负责安全教育、危险源辨识、临电脚手架检查和安全资料管理。"],
            ["材料员", "负责材料计划、进场验收、保管发放和周转材料管理。"],
            ["资料员", "负责施工技术资料、试验报告、验收记录和竣工资料归档。"],
        ],
        widths=[1800, 7560],
    )
    add_heading(doc, "3.3 施工程序与流水段划分", 2)
    add_para(doc, "工程按基础、主体结构、砌体围护、屋面防水、装饰装修、室外及竣工验收六个阶段组织。主体结构按建筑平面东西向划分为两个施工段，每层形成钢筋绑扎、模板安装、混凝土浇筑和养护流水作业。砌体和安装预留预埋在主体结构完成二层后适时插入，装饰工程在主体结构验收后分层展开。")
    add_numbered(
        doc,
        [
            "施工准备：完成临建、临水临电、测量控制网、图纸会审、材料供应计划和机械进场验收。",
            "基础施工：测量放线、土方开挖、基槽验收、垫层、钢筋模板、基础混凝土和回填土。",
            "主体施工：柱、梁、板、楼梯模板钢筋混凝土流水施工，安装预留预埋同步配合。",
            "围护与屋面：砌体、构造柱、圈梁、屋面找坡、保温、防水和保护层施工。",
            "装饰装修：内外墙抹灰、楼地面、门窗、涂饰、外墙真石漆及细部收口。",
            "收尾验收：系统调试、质量整改、资料整理、竣工验收和移交保修。",
        ],
    )
    add_heading(doc, "3.4 机械设备进场计划", 2)
    add_caption(doc, "表3.1 机械设备进场计划")
    add_table(
        doc,
        ["序号", "设备名称", "规格型号", "数量", "使用部位", "进场时间"],
        [
            ["1", "塔式起重机", "QTZ63", "1台", "主体垂直运输", "开工第20天"],
            ["2", "施工升降机", "SC200/200", "1台", "砌体装饰运输", "主体三层后"],
            ["3", "挖掘机", "PC200", "2台", "土方开挖", "开工第3天"],
            ["4", "自卸汽车", "15t", "6辆", "土方外运", "开工第3天"],
            ["5", "钢筋切断机", "GQ40", "2台", "钢筋加工", "开工第10天"],
            ["6", "钢筋弯曲机", "GW40", "2台", "钢筋加工", "开工第10天"],
            ["7", "木工圆盘锯", "MJ104", "2台", "模板加工", "开工第12天"],
            ["8", "混凝土振捣器", "ZX50", "8台", "混凝土浇筑", "基础施工前"],
            ["9", "砂浆搅拌机", "UJ325", "2台", "砌筑抹灰", "砌体施工前"],
            ["10", "电焊机", "BX1-400", "3台", "钢筋及预埋件", "开工第10天"],
            ["11", "水泵", "QY25", "4台", "基坑排水", "土方开挖前"],
        ],
        widths=[700, 1700, 1500, 900, 2600, 1960],
    )
    add_heading(doc, "3.5 试验和检测仪器设备", 2)
    add_caption(doc, "表3.2 试验和检测仪器设备")
    add_table(
        doc,
        ["序号", "仪器、设备名称", "规格型号", "数量", "用途"],
        [
            ["1", "全站仪", "NTS-332R", "1台", "轴线测设、坐标控制"],
            ["2", "水准仪", "DS3", "2台", "标高测量"],
            ["3", "钢卷尺", "50m", "4把", "距离测量"],
            ["4", "坍落度筒", "标准型", "2套", "混凝土坍落度检测"],
            ["5", "混凝土试模", "150mm", "18组", "混凝土试块制作"],
            ["6", "砂浆试模", "70.7mm", "12组", "砂浆试块制作"],
            ["7", "靠尺塞尺", "2m", "3套", "抹灰、地面平整度检测"],
            ["8", "回弹仪", "ZC3-A", "1台", "混凝土实体强度检测"],
            ["9", "接地电阻仪", "ZC-8", "1台", "临电接地检测"],
            ["10", "游标卡尺", "0-200mm", "2把", "钢筋直径及构件尺寸检测"],
        ],
        widths=[700, 2300, 1600, 900, 3860],
    )
    add_heading(doc, "3.6 劳动力计划与分包计划", 2)
    add_para(doc, "劳动力按施工进度分阶段投入，高峰期主要集中在主体结构和装饰装修阶段。劳动力组织应保持相对均衡，并根据关键线路工序适时调整班组数量。")
    add_caption(doc, "表3.3 劳动力计划表")
    add_table(
        doc,
        ["工种", "准备", "基础", "主体", "砌体", "装饰", "收尾"],
        [
            ["木工", "8", "22", "45", "10", "8", "4"],
            ["钢筋工", "6", "28", "38", "6", "4", "2"],
            ["混凝土工", "4", "18", "24", "6", "6", "2"],
            ["瓦工", "2", "8", "12", "36", "22", "8"],
            ["抹灰工", "0", "0", "0", "12", "42", "10"],
            ["架子工", "4", "8", "16", "12", "10", "6"],
            ["防水工", "0", "0", "0", "4", "12", "4"],
            ["普工", "12", "28", "36", "32", "30", "16"],
            ["合计", "36", "112", "171", "118", "134", "52"],
        ],
        widths=[1600, 1290, 1290, 1290, 1290, 1290, 1310],
    )
    add_para(doc, "拟分包工程包括防水工程、外墙保温及真石漆工程、门窗工程和脚手架搭拆工程。分包单位应具备相应资质，进场前完成资格审查、安全教育和技术交底，施工质量纳入总承包单位统一管理。")

    add_heading(doc, "4 施工总平面布置", 1)
    add_para(doc, "施工总平面布置以减少二次搬运、保证交通顺畅、满足安全文明施工和消防要求为原则。现场设置封闭围挡，北侧布置主出入口和洗车设施，西北角布置办公生活区，建筑物西侧布置钢筋加工棚和模板加工棚，东侧设周转材料堆场，塔吊布置在建筑物南侧偏中位置，覆盖主体施工和主要材料堆场。")
    add_bullets(
        doc,
        [
            "场内道路采用200mm厚C25混凝土硬化，主运输道路宽度不小于4.0m，转弯处满足材料车辆通行。",
            "钢筋、模板、砌块、水泥砂浆等材料堆场靠近施工电梯和塔吊覆盖区域，堆放分区标识清晰。",
            "临时用电采用三级配电、二级保护，电缆沿围挡或架空线路敷设，穿越道路处加设保护套管。",
            "现场排水采用明沟与沉淀池结合，洗车废水沉淀后循环使用，不得直接排入市政管网。",
            "办公区、生活区、加工区和施工区分隔设置，消防器材按规定配置，动火作业执行审批制度。",
        ],
    )
    add_para(doc, "施工总平面布置示意详见附件2。正式成果提交时，可按任务书要求将平面图绘制为A3图幅，补充图框、图例、指北针、绘图人和审核人信息后折叠装订。")

    add_heading(doc, "5 施工进度计划", 1)
    add_para(doc, "本工程计划总工期300日历天。施工进度安排综合考虑校园环境、季节影响、结构流水施工和装饰穿插作业。关键线路为施工准备、土方及基础、主体结构、砌体及屋面、内外装饰、室外收尾、竣工验收。")
    add_caption(doc, "表5.1 关键工期控制节点")
    add_table(
        doc,
        ["节点", "计划完成时间", "控制要求"],
        [
            ["开工准备完成", "第10天", "临建、临电、临水、测量控制网及材料计划完成"],
            ["基础结构完成", "第60天", "基础验收合格，具备主体施工条件"],
            ["主体结构封顶", "第150天", "五层主体结构完成并通过主体结构验收"],
            ["砌体及屋面完成", "第195天", "砌体、构造柱、屋面防水完成"],
            ["装饰装修完成", "第270天", "内外装饰、门窗、楼地面基本完成"],
            ["竣工验收", "第300天", "资料齐全，现场清理完成，通过竣工验收"],
        ],
        widths=[2200, 1900, 5260],
    )
    add_para(doc, "进度控制采用月计划、周计划和日碰头制度。关键工作出现延误时，优先通过增加作业面、调整流水节拍、组织平行施工和优化材料供应进行纠偏。施工进度计划横道图详见附件1。")

    add_heading(doc, "6 主要分部分项工程施工方案", 1)
    sections = [
        ("6.1 土方施工方案", "土方开挖前复核坐标和标高，做好地下管线调查、围挡封闭、施工道路和排水设施。基坑采用机械开挖、人工清底方式，开挖至设计标高以上200mm时由人工清底，严禁扰动持力层。基坑周边设置防护栏杆和警示标识，雨季施工时设置集水坑和水泵，保证基坑内无积水。土方外运车辆出场前冲洗，运输过程覆盖，减少扬尘污染。基槽完成后及时组织验槽，验槽合格后浇筑垫层。"),
        ("6.2 地基与基础施工方案", "基础施工顺序为测量放线、垫层施工、钢筋绑扎、模板安装、隐蔽验收、混凝土浇筑、养护和回填。钢筋进场应有质量证明文件并按批复试，钢筋绑扎位置、保护层厚度、锚固长度和搭接长度应符合设计及规范要求。基础模板应具有足够强度、刚度和稳定性，混凝土浇筑前清理模板内杂物并湿润。混凝土采用商品混凝土，分层振捣密实，浇筑完成后覆盖保湿养护不少于7d。"),
        ("6.3 主体结构施工方案", "主体结构采用现浇钢筋混凝土框架体系，按楼层和施工段组织流水施工。模板支撑体系编制专项方案并经验收合格后使用，立杆间距、扫地杆、剪刀撑和可调托撑设置符合规范要求。梁板柱钢筋绑扎前进行翻样和技术交底，节点核心区箍筋、梁柱锚固和板负筋位置作为质量控制重点。混凝土浇筑按先柱墙后梁板、先远后近原则进行，振捣密实并控制施工缝留设位置。"),
        ("6.4 脚手架施工方案", "外脚手架采用落地式双排扣件式钢管脚手架，搭设前夯实基础并设置垫板和排水措施。脚手架纵距、横距、步距、连墙件、剪刀撑和安全网按专项方案实施。搭设、使用和拆除阶段均由专职安全员检查，未经项目技术负责人和安全负责人验收不得投入使用。拆除作业遵循先搭后拆、后搭先拆原则，设置警戒区并安排专人监护。"),
        ("6.5 砌筑工程施工方案", "砌体施工前完成主体结构验收、弹线、植筋和构造柱钢筋绑扎。砌块提前适度湿润，砂浆按配合比计量拌制，砌筑时控制灰缝厚度、墙体垂直度和平整度。填充墙顶部采用斜砌或柔性连接措施，待下部墙体沉实后再施工。门窗洞口、构造柱、圈梁、过梁和拉结筋设置应符合设计要求。"),
        ("6.6 抹灰与楼地面施工方案", "抹灰前清理基层、浇水湿润、做灰饼冲筋，不同材料交接处铺设耐碱网格布或钢丝网。抹灰分层进行，每层厚度和间隔时间符合规范要求，阴阳角方正、表面平整。楼地面施工前进行基层清理和标高控制，水磨石或地砖面层应先排版后铺贴，卫生间等有水房间按规定做防水和蓄水试验。"),
        ("6.7 门窗与装饰工程施工方案", "门窗框安装前复核洞口尺寸、标高和垂直度，采用可靠连接件固定，框与墙体间缝隙填嵌密实并打密封胶。内墙涂饰施工前基层含水率和腻子质量应符合要求，涂料涂刷均匀、颜色一致、无明显流坠。外墙真石漆施工应在基层平整、干燥、分格缝完成后进行，喷涂厚度均匀，成品保护到位。"),
        ("6.8 屋面与防水施工方案", "屋面施工顺序为基层处理、找坡层、保温层、找平层、防水附加层、防水卷材、保护层。防水卷材施工前基层应平整、干燥、阴阳角做成圆弧，女儿墙、落水口、管根等节点增设附加层。卷材搭接宽度、热熔粘结质量和收头固定应符合规范要求。防水完成后进行淋水或蓄水试验，合格后施工保护层。"),
    ]
    for heading, text in sections:
        add_heading(doc, heading, 2)
        add_para(doc, text)
        add_para(doc, "质量验收重点包括原材料复试、隐蔽工程验收、检验批质量记录、实测实量和观感质量。施工前应进行班组技术交底，施工中执行样板引路、过程旁站和三检制度，发现偏差及时整改并形成闭合记录。")

    add_heading(doc, "7 主要管理和技术组织措施", 1)
    measures = [
        ("7.1 进度目标控制措施", "建立项目总进度计划、月计划、周计划和日作业计划体系。项目经理每周组织生产例会，检查关键线路工作完成情况，分析材料、劳动力、机械和技术问题。对影响工期的图纸、材料和分包问题提前预警，必要时通过增加班组、延长作业时间、调整流水节拍和优化施工顺序进行纠偏。"),
        ("7.2 质量目标控制措施", "实行项目经理负责、技术负责人主控、质量员专检、班组自检互检的质量保证体系。坚持图纸会审、方案审批、技术交底、样板引路、材料复试、隐蔽验收和检验批验收制度。钢筋、模板、混凝土、防水、抹灰和外墙装饰等关键工序设置质量控制点。"),
        ("7.3 安全目标控制措施", "建立安全生产责任制和危险源清单，开展三级安全教育和班前安全活动。对深基坑、脚手架、模板支撑、临时用电、塔吊和高处作业等危险性较大的分部分项工程编制专项方案并组织验收。现场设置安全通道、防护棚、临边洞口防护和消防设施。"),
        ("7.4 成本目标控制措施", "成本控制从施工方案、材料采购、周转材料利用、机械台班、劳动力效率和现场签证等方面入手。材料实行计划采购、进场验收、限额领料和余料回收；模板、脚手架和机械设备提高周转效率；设计变更和现场签证及时办理，保证成本数据真实完整。"),
        ("7.5 季节性施工措施", "雨季施工前疏通排水系统，材料堆场垫高覆盖，基坑和道路设置排水沟、集水坑和水泵。高温季节调整作息时间，混凝土浇筑避开高温时段并加强保湿养护。冬期低温施工时根据气温采取保温、防冻和测温措施，保证砂浆、混凝土质量。"),
        ("7.6 文明施工与环境保护措施", "现场实行封闭管理，主要道路硬化并定时洒水降尘，出入口设置车辆冲洗设施。易扬尘材料覆盖堆放，建筑垃圾分类收集并及时清运。合理安排噪声较大的作业时间，避免影响学校和周边居民。施工现场标识统一、材料堆放整齐、生活区卫生整洁。"),
        ("7.7 成品保护措施", "建立成品保护责任区和交接制度。钢筋绑扎完成后设置马凳和通道板，防止踩踏；混凝土浇筑后按规定养护并控制上人时间；门窗、地面、墙面、外墙涂料和防水层施工完成后采取覆盖、封闭和专人看护措施，避免后续工序污染破坏。"),
        ("7.8 新技术、新工艺应用", "本工程拟采用钢筋集中加工配送、模板早拆体系、预拌砂浆、二维码材料验收台账、移动端质量安全巡检和节水降尘循环利用等措施，提高施工效率和管理精细化水平。"),
        ("7.9 治安保卫与工程保修", "施工现场实行门卫登记和实名制管理，材料库房、危险品仓库和夜间值班由专人负责。工程竣工后建立质量回访和保修台账，按合同和国家规定履行保修义务，对用户反馈问题及时响应、维修和记录。"),
    ]
    for h, t in measures:
        add_heading(doc, h, 2)
        add_para(doc, t)

    add_heading(doc, "8 主要技术经济指标", 1)
    add_para(doc, "主要技术经济指标用于评价施工组织设计的合理性和施工管理效果。项目实施过程中应结合实际工程量、合同价和资源消耗进行动态复核。")
    add_table(
        doc,
        ["指标类别", "控制指标", "目标值"],
        [
            ["质量指标", "单位工程一次验收合格率", "100%"],
            ["工期指标", "合同工期履约率", "100%，总工期300日历天"],
            ["安全指标", "较大及以上安全事故", "0"],
            ["文明施工", "扬尘和噪声投诉", "力争为0，满足地方管理要求"],
            ["材料管理", "主要材料损耗率", "钢筋不大于2%，商品混凝土按计划控制"],
            ["机械管理", "主要机械完好率", "不低于90%"],
            ["资料管理", "技术资料同步率", "不低于95%"],
            ["环保指标", "建筑垃圾分类清运率", "100%"],
        ],
        widths=[2000, 3700, 3660],
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "《建筑工程施工质量验收统一标准》GB 50300-2013.",
        "《混凝土结构工程施工质量验收规范》GB 50204-2015.",
        "《砌体结构工程施工质量验收规范》GB 50203-2011.",
        "《建筑施工安全检查标准》JGJ 59-2011.",
        "《建筑施工扣件式钢管脚手架安全技术规范》JGJ 130-2011.",
        "《建筑施工模板安全技术规范》JGJ 162-2008.",
        "《建筑施工组织设计规范》GB/T 50502-2009.",
        "《建筑施工现场环境与卫生标准》JGJ 146-2013.",
    ]
    add_numbered(doc, refs)


def appendix(doc):
    sec = doc.add_section(WD_ORIENT.LANDSCAPE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    add_heading(doc, "附件1 施工进度计划横道图", 1)
    headers = ["施工内容", "1月", "2月", "3月", "4月", "5月", "6月", "7月", "8月", "9月", "10月"]
    rows = [
        ["施工准备", "■■", "", "", "", "", "", "", "", "", ""],
        ["土方及基础", "■", "■■", "", "", "", "", "", "", "", ""],
        ["主体一至三层", "", "■", "■■", "", "", "", "", "", "", ""],
        ["主体四至五层及屋面结构", "", "", "■", "■■", "", "", "", "", "", ""],
        ["砌体及二次结构", "", "", "", "■", "■■", "", "", "", "", ""],
        ["屋面防水", "", "", "", "", "■■", "", "", "", "", ""],
        ["内外墙抹灰", "", "", "", "", "■", "■■", "", "", "", ""],
        ["楼地面与门窗", "", "", "", "", "", "■", "■■", "", "", ""],
        ["外墙装饰", "", "", "", "", "", "", "■■", "■", "", ""],
        ["室外及竣工验收", "", "", "", "", "", "", "", "■", "■■", "■"],
    ]
    add_table(doc, headers, rows, widths=[2200] + [900] * 10)
    add_para(doc, "说明：黑色方块表示该分项工程主要施工时段。正式提交时可按A3图幅进一步采用双代号时标网络计划或专业进度计划软件绘制，并标明关键线路。", first_line=False)

    add_heading(doc, "附件2 施工现场总平面布置示意图", 1)
    site_plan = draw_site_plan(SITE_PLAN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(site_plan), width=Inches(10.15))
    add_para(doc, "说明：本图为课程设计阶段施工总平面布置示意，已表达主出入口、临时道路、拟建建筑物、塔吊、加工棚、材料堆场、办公生活区、临水临电、消防水源及沉淀池等主要要素。正式A3图幅可在此基础上按比例深化。", first_line=False)


def main():
    doc = Document()
    configure_doc(doc)
    cover(doc)
    toc(doc)
    body(doc)
    appendix(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
